from __future__ import annotations

import json
import random
import re
import urllib.request
from collections import defaultdict
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from openpyxl import load_workbook


ROOT_DIR = Path(__file__).resolve().parent.parent
ENV_PATH = ROOT_DIR / ".env"
OUTPUT_DIR = ROOT_DIR / "boost outputs"
DEFAULT_DURATION_MINUTES = 20
TASKS_PER_STUDENT = 8
GRAMMAR_TASK_NUMBERS = list(range(20, 29))
EXCLUDED_STUDENT_KEYS = {"дугинец", "выступец_дарья"}
TOPIC_LABELS = {
    20: "видовременные формы / базовые формы глагола",
    21: "пассивный залог / длительные формы / согласование времен",
    22: "модальные формы / условные конструкции / формы глагола",
    23: "степени сравнения / числительные / местоимения",
    24: "местоимения / притяжательные формы / личные формы",
    25: "множественное число / исключения / существительные",
    26: "числительные / порядковые формы",
    27: "пассивный залог / причастия / формы глагола",
    28: "сравнительные конструкции / лексико-грамматические формы",
}


@dataclass
class BankTask:
    variant_number: int
    task_number: int
    order_index: int
    topic: str
    cue_word: str
    prompt_text: str
    answer: str


@dataclass
class StudentBoostTask:
    title: str
    full_name: str
    student_key: str
    group_name: str
    weak_topics: list[str]
    selected_tasks: list[BankTask]


def main() -> None:
    env = load_env(ENV_PATH)
    xlsx_path = resolve_input_path(env)
    tasks = fetch_grammar_bank_tasks(env)
    students = analyze_students(xlsx_path)
    generated = build_student_variants(students, tasks)

    timestamp = file_timestamp()
    output_dir = OUTPUT_DIR / f"grammar-boost-{timestamp}"
    output_dir.mkdir(parents=True, exist_ok=True)

    create_student_documents(generated, output_dir)
    create_answer_keys_document(generated, output_dir / "teacher-answer-keys.docx")
    write_manifest(generated, output_dir / "manifest.json")

    print(f"Generated {len(generated)} student grammar boost files in {output_dir}")
    print(f"Teacher keys: {output_dir / 'teacher-answer-keys.docx'}")


def load_env(path: Path) -> dict[str, str]:
    env: dict[str, str] = {}
    if not path.exists():
      return env
    for line in path.read_text("utf-8").splitlines():
        stripped = line.strip()
        if not stripped or stripped.startswith("#") or "=" not in stripped:
            continue
        key, value = stripped.split("=", 1)
        env[key.strip()] = value.strip()
    return env


def resolve_input_path(env: dict[str, str]) -> Path:
    configured = env.get("PB_XLSX_FILE", "")
    if configured:
        candidate = (ROOT_DIR / configured).resolve()
        if candidate.exists():
            return candidate

    files = sorted(ROOT_DIR.glob("*.xlsx"), key=lambda item: item.stat().st_mtime, reverse=True)
    if files:
        return files[0]

    fallback = ROOT_DIR / "Англ 9 Север-2.xlsx"
    if fallback.exists():
        return fallback
    raise RuntimeError("No xlsx file found in the project root.")


def fetch_grammar_bank_tasks(env: dict[str, str]) -> list[BankTask]:
    base_url = env.get("PB_URL", "http://127.0.0.1:8091").rstrip("/")
    email = env.get("PB_ADMIN_EMAIL", "")
    password = env.get("PB_ADMIN_PASSWORD", "")

    token = authenticate(base_url, email, password)
    records = fetch_all(base_url, "oge_grammar_tasks", token)
    tasks = [
        BankTask(
            variant_number=int(record["variantNumber"]),
            task_number=int(record["taskNumber"]),
            order_index=int(record["orderIndex"]),
            topic=str(record["topic"]),
            cue_word=str(record["cueWord"]),
            prompt_text=str(record["promptText"]),
            answer=str(record["answer"]),
        )
        for record in records
    ]
    tasks.sort(key=lambda item: (item.task_number, item.variant_number, item.order_index))
    return tasks


def authenticate(base_url: str, email: str, password: str) -> str:
    if not email or not password:
        raise RuntimeError("PB_ADMIN_EMAIL and PB_ADMIN_PASSWORD are required.")
    data = json.dumps({"identity": email, "password": password}).encode("utf-8")
    request = urllib.request.Request(
        f"{base_url}/api/collections/_superusers/auth-with-password",
        data=data,
        method="POST",
        headers={"Content-Type": "application/json", "Accept": "application/json"},
    )
    with urllib.request.urlopen(request) as response:
        payload = json.loads(response.read().decode("utf-8"))
    return str(payload["token"])


def fetch_all(base_url: str, collection: str, token: str) -> list[dict[str, object]]:
    page = 1
    items: list[dict[str, object]] = []
    while True:
        request = urllib.request.Request(
            f"{base_url}/api/collections/{collection}/records?page={page}&perPage=200&sort=variantNumber,orderIndex",
            headers={"Authorization": f"Bearer {token}", "Accept": "application/json"},
        )
        with urllib.request.urlopen(request) as response:
            payload = json.loads(response.read().decode("utf-8"))
        items.extend(payload.get("items", []))
        if page >= payload.get("totalPages", 1):
            break
        page += 1
    return items


def analyze_students(xlsx_path: Path) -> list[dict[str, object]]:
    workbook = load_workbook(xlsx_path, data_only=True)
    sheet_names = [name for name in workbook.sheetnames if not name.lower().startswith("шаблон")]
    student_map: dict[str, dict[str, object]] = {}

    for sheet_name in reversed(sheet_names):
        variant_number = extract_variant_number(sheet_name)
        if variant_number <= 0:
            continue

        ws = workbook[sheet_name]
        rows = list(ws.iter_rows(values_only=True))
        maximums = list(rows[4]) if len(rows) > 4 else []
        prev_num: int | None = None
        group_index = 1

        for row in rows[5:]:
            values = list(row)
            raw_num = numeric(values[0] if len(values) > 0 else None)
            raw_name = string_value(values[1] if len(values) > 1 else None)

            if not raw_name:
                continue
            if "среднее значение" in raw_name.lower():
                break

            if prev_num is not None and raw_num < prev_num and raw_num <= 2:
                group_index += 1
            prev_num = raw_num

            full_name = raw_name.strip()
            student_key = normalize_name(full_name)
            if is_excluded_student(student_key, full_name):
                continue

            row_slice = values[2:37]
            if not any(string_value(value) for value in row_slice):
                continue

            student = student_map.setdefault(
                student_key,
                {
                    "full_name": full_name,
                    "student_key": student_key,
                    "group_name": f"9.{group_index}",
                    "task_stats": defaultdict(lambda: {"sum": 0.0, "count": 0, "solved_variants": set()}),
                },
            )

            for task_number in GRAMMAR_TASK_NUMBERS:
                column_index = task_number + 1
                raw_value = string_value(values[column_index] if column_index < len(values) else None)
                if not raw_value:
                    continue
                max_score = numeric(maximums[column_index] if column_index < len(maximums) else None)
                if max_score <= 0:
                    continue
                score = numeric(values[column_index] if column_index < len(values) else None)
                percent = round((score / max_score) * 100, 1)
                task_stat = student["task_stats"][task_number]
                task_stat["sum"] += percent
                task_stat["count"] += 1
                task_stat["solved_variants"].add(variant_number)

    return sorted(student_map.values(), key=lambda item: (str(item["group_name"]), str(item["full_name"])))


def build_student_variants(students: list[dict[str, object]], tasks: list[BankTask]) -> list[StudentBoostTask]:
    tasks_by_number: dict[int, list[BankTask]] = defaultdict(list)
    for task in tasks:
        tasks_by_number[task.task_number].append(task)

    generated: list[StudentBoostTask] = []
    for student in students:
        ranked_numbers = rank_weak_tasks(student["task_stats"])
        selected = select_bank_tasks(str(student["student_key"]), ranked_numbers, tasks_by_number)
        weak_topics = [TOPIC_LABELS[number] for number in ranked_numbers[:4]]
        generated.append(
            StudentBoostTask(
                title="Grammar Boost Variant",
                full_name=str(student["full_name"]),
                student_key=str(student["student_key"]),
                group_name=str(student["group_name"]),
                weak_topics=weak_topics,
                selected_tasks=selected,
            )
        )

    return generated


def rank_weak_tasks(task_stats: dict[int, dict[str, object]]) -> list[int]:
    ranked: list[tuple[float, int, int]] = []
    for task_number in GRAMMAR_TASK_NUMBERS:
        stat = task_stats.get(task_number)
        if not stat or int(stat["count"]) == 0:
            ranked.append((100.0, 0, task_number))
            continue
        average_percent = float(stat["sum"]) / int(stat["count"])
        ranked.append((average_percent, -int(stat["count"]), task_number))
    ranked.sort()
    return [task_number for _, _, task_number in ranked]


def select_bank_tasks(
    student_key: str,
    ranked_numbers: list[int],
    tasks_by_number: dict[int, list[BankTask]],
) -> list[BankTask]:
    selected: list[BankTask] = []
    used_keys: set[tuple[int, int]] = set()
    rng = random.Random(student_key)

    priority_numbers = ranked_numbers[:4]
    for task_number in priority_numbers:
        pool = list(tasks_by_number.get(task_number, []))
        if not pool:
            continue
        rng.shuffle(pool)
        for task in pool[:2]:
            key = (task.variant_number, task.task_number)
            if key in used_keys:
                continue
            selected.append(task)
            used_keys.add(key)
            if len(selected) >= TASKS_PER_STUDENT:
                return sort_selected_tasks(selected)

    if len(selected) < TASKS_PER_STUDENT:
        for task_number in ranked_numbers[4:]:
            pool = list(tasks_by_number.get(task_number, []))
            rng.shuffle(pool)
            for task in pool:
                key = (task.variant_number, task.task_number)
                if key in used_keys:
                    continue
                selected.append(task)
                used_keys.add(key)
                if len(selected) >= TASKS_PER_STUDENT:
                    return sort_selected_tasks(selected)

    return sort_selected_tasks(selected[:TASKS_PER_STUDENT])


def sort_selected_tasks(tasks: list[BankTask]) -> list[BankTask]:
    return sorted(tasks, key=lambda item: (item.task_number, item.variant_number, item.order_index))


def create_student_documents(items: list[StudentBoostTask], output_dir: Path) -> None:
    for item in items:
        document = Document()
        set_default_style(document)

        title = document.add_heading(f"{item.full_name} — Grammar Boost", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        meta = document.add_paragraph()
        meta.add_run(f"Группа {item.group_name} · до {DEFAULT_DURATION_MINUTES} минут · 8 заданий").bold = True

        topics = document.add_paragraph()
        topics.add_run("Темы отработки: ").bold = True
        topics.add_run("; ".join(item.weak_topics) if item.weak_topics else "грамматика")

        document.add_paragraph("")

        for index, task in enumerate(item.selected_tasks, start=1):
            paragraph = document.add_paragraph(style=None)
            paragraph.paragraph_format.space_after = Pt(6)
            run = paragraph.add_run(f"{index}. ")
            run.bold = True
            paragraph.add_run(task.prompt_text)
            cue = document.add_paragraph()
            cue.paragraph_format.left_indent = Pt(18)
            cue.add_run(task.cue_word).bold = True

        filename = sanitize_filename(item.full_name) + ".docx"
        document.save(output_dir / filename)


def create_answer_keys_document(items: list[StudentBoostTask], output_path: Path) -> None:
    document = Document()
    set_default_style(document)
    heading = document.add_heading("Grammar Boost — Answer Keys", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for item in items:
        document.add_heading(item.full_name, level=2)
        meta = document.add_paragraph()
        meta.add_run(f"Группа {item.group_name}").italic = True
        for index, task in enumerate(item.selected_tasks, start=1):
            paragraph = document.add_paragraph(style=None)
            paragraph.add_run(f"{index}. ").bold = True
            paragraph.add_run(f"{task.cue_word} → {task.answer} ")
            source = paragraph.add_run(f"(variant {task.variant_number}, task {task.task_number})")
            source.italic = True

    document.save(output_path)


def write_manifest(items: list[StudentBoostTask], output_path: Path) -> None:
    payload = [
        {
            "fullName": item.full_name,
            "groupName": item.group_name,
            "weakTopics": item.weak_topics,
            "tasks": [
                {
                    "variantNumber": task.variant_number,
                    "taskNumber": task.task_number,
                    "cueWord": task.cue_word,
                    "answer": task.answer,
                    "topic": task.topic,
                }
                for task in item.selected_tasks
            ],
        }
        for item in items
    ]
    output_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), "utf-8")


def set_default_style(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)


def extract_variant_number(sheet_name: str) -> int:
    match = re.search(r"(?:^|\s)(\d+)\s*TEST\b", sheet_name, re.I)
    return int(match.group(1)) if match else 0


def numeric(value: object) -> float:
    if value is None:
        return 0.0
    try:
        return float(str(value).replace(",", "."))
    except ValueError:
        return 0.0


def string_value(value: object) -> str:
    return str(value or "").strip()


def normalize_name(value: str) -> str:
    return (
        value.strip()
        .lower()
        .split("(")[0]
        .strip()
        .replace("ё", "е")
        .replace(" ", "-")
    )


def is_excluded_student(student_key: str, full_name: str) -> bool:
    normalized_full_name = full_name.lower().replace("ё", "е")
    return (
        student_key in EXCLUDED_STUDENT_KEYS
        or "дугинец" in normalized_full_name
        or "выступец дарья" in normalized_full_name
    )


def sanitize_filename(value: str) -> str:
    return re.sub(r"[^\w\-а-яА-ЯёЁ ]+", "", value).strip().replace("/", "-")


def file_timestamp() -> str:
    from datetime import datetime

    return datetime.now().strftime("%Y-%m-%d_%H-%M")


if __name__ == "__main__":
    main()
